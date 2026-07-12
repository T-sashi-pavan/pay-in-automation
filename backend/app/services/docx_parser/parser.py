from io import BytesIO
from typing import List, Dict, Any, Optional
from docx import Document

from backend.app.services.excel_parser.parser import ExcelParserService, is_rule_effectively_empty


class DocxParserService:
    """
    Parses commission-grid tables out of a Word document. Word tables are
    inherently flat (one row = one rule), so each table is routed straight
    through ExcelParserService.parse_table — the same column-classification,
    normalization, and slab-merge pipeline Excel FLAT sheets use.
    """

    def __init__(self):
        self.excel_parser = ExcelParserService()

    def parse_document(self, file_bytes: bytes, filename: Optional[str] = None) -> List[Dict[str, Any]]:
        doc = Document(BytesIO(file_bytes))

        company = "Unknown"
        if filename:
            fn_lower = filename.lower()
            if "tata" in fn_lower:
                company = "Tata"
            elif "shriram" in fn_lower:
                company = "Shriram"
            elif "chola" in fn_lower:
                company = "Cholamandalam"
            elif "digit" in fn_lower:
                company = "Digit"
            elif "oriental" in fn_lower:
                company = "Oriental"

        all_parsed_rules: List[Dict[str, Any]] = []
        existing_keys: set = set()

        print(f"\n[STARTING WORD DOCUMENT PARSING] '{filename}' — {len(doc.tables)} table(s) detected")

        for table_idx, table in enumerate(doc.tables):
            table_rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            table_rows = [r for r in table_rows if any(c for c in r)]
            if len(table_rows) < 2:
                continue

            headers = table_rows[0]
            data_rows = table_rows[1:]
            sheet_name = f"Table {table_idx + 1}"

            count = self.excel_parser.parse_table(
                headers, data_rows, sheet_name, filename, company, existing_keys, all_parsed_rules
            )
            print(f"  [TABLE COMPLETED] Extracted {count} rules from '{sheet_name}'.")

        final_rules = self.excel_parser._group_and_merge_rules(all_parsed_rules)
        final_rules = [r for r in final_rules if not is_rule_effectively_empty(r)]
        print(f"[WORD DOCUMENT PARSING COMPLETED] Grouped {len(all_parsed_rules)} rules into {len(final_rules)} merged rules.")
        return final_rules
